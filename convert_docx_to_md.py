#!/usr/bin/env python3
from docx import Document
import sys
import os

def docx_to_markdown(docx_path, md_path):
    doc = Document(docx_path)
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append('')
            continue
        
        # 检查段落样式
        style = para.style.name.lower()
        
        if 'heading 1' in style or '标题 1' in style:
            md_lines.append(f'# {text}')
        elif 'heading 2' in style or '标题 2' in style:
            md_lines.append(f'## {text}')
        elif 'heading 3' in style or '标题 3' in style:
            md_lines.append(f'### {text}')
        elif 'heading 4' in style or '标题 4' in style:
            md_lines.append(f'#### {text}')
        elif 'heading 5' in style or '标题 5' in style:
            md_lines.append(f'##### {text}')
        elif 'heading 6' in style or '标题 6' in style:
            md_lines.append(f'###### {text}')
        elif 'list' in style or '列表' in style:
            # 尝试判断是否是有序列表
            if para.text and (para.text[0].isdigit() or 'number' in style.lower()):
                md_lines.append(f'1. {text}')
            else:
                md_lines.append(f'- {text}')
        else:
            md_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        md_lines.append('')
        # 表头
        header_cells = table.rows[0].cells
        header = '| ' + ' | '.join(cell.text.strip() for cell in header_cells) + ' |'
        md_lines.append(header)
        # 分隔线
        separator = '| ' + ' | '.join('---' for _ in header_cells) + ' |'
        md_lines.append(separator)
        # 数据行
        for row in table.rows[1:]:
            row_cells = row.cells
            row_text = '| ' + ' | '.join(cell.text.strip() for cell in row_cells) + ' |'
            md_lines.append(row_text)
        md_lines.append('')
    
    md_content = '\n'.join(md_lines)
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    print(f'转换完成: {md_path}')

if __name__ == '__main__':
    docx_path = '/Users/dada/Documents/GitHub/Project-Ark/Docs/静默方舟.docx'
    md_path = '/Users/dada/Documents/GitHub/Project-Ark/Docs/静默方舟.md'
    
    docx_to_markdown(docx_path, md_path)
