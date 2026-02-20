#!/usr/bin/env python3
from docx import Document
import sys

def inspect_docx(docx_path):
    doc = Document(docx_path)
    print(f"=== 检查文件: {docx_path} ===")
    print(f"段落总数: {len(doc.paragraphs)}")
    print(f"表格总数: {len(doc.tables)}")
    print("\n=== 前50个段落 ===")
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        print(f"{i+1}. [{para.style.name}] {repr(text)}")
    
    if len(doc.paragraphs) > 50:
        print(f"... 还有 {len(doc.paragraphs)-50} 个段落")
    
    print("\n=== 表格 ===")
    for i, table in enumerate(doc.tables):
        print(f"表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")

if __name__ == '__main__':
    inspect_docx('/Users/dada/Documents/GitHub/Project-Ark/Docs/GDD/示巴星 (鸣钟者)/示巴星 V1.docx')
    print("\n" + "="*80 + "\n")
    inspect_docx('/Users/dada/Documents/GitHub/Project-Ark/Docs/GDD/示巴星 (鸣钟者)/示巴星 V2.docx')
