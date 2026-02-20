#!/usr/bin/env python3
from docx import Document
from docx.text.paragraph import Paragraph
import sys
import os

def get_full_text(para):
    full_text = []
    for run in para.runs:
        if run.text:
            full_text.append(run.text)
    return ''.join(full_text).strip()

def docx_to_markdown(docx_path, md_path):
    doc = Document(docx_path)
    md_lines = []
    
    for para in doc.paragraphs:
        text = get_full_text(para)
        if not text:
            md_lines.append('')
            continue
        
        style = para.style.name.lower()
        
        if 'heading 1' in style or '标题 1' in style:
            md_lines.append(f'# {text}')
        elif 'heading 2' in style or '标题 2' in style:
            md_lines.append(f'## {text}')
        elif 'heading 3' in style or '标题 3' in style:
            md_lines.append(f'### {text}')
        elif 'heading 4' in style or '标题 4' in style:
            md_lines.append(f'#### {text}')
        elif 'heading 5' in style or '标题 5' in style:
            md_lines.append(f'##### {text}')
        elif 'heading 6' in style or '标题 6' in style:
            md_lines.append(f'###### {text}')
        elif 'list' in style or '列表' in style:
            if text and (text[0].isdigit() or 'number' in style.lower()):
                md_lines.append(f'1. {text}')
            else:
                md_lines.append(f'- {text}')
        else:
            md_lines.append(text)
    
    for table in doc.tables:
        md_lines.append('')
        header_cells = table.rows[0].cells
        header_texts = [get_full_text(cell.paragraphs[0]) for cell in header_cells]
        header = '| ' + ' | '.join(header_texts) + ' |'
        md_lines.append(header)
        separator = '| ' + ' | '.join('---' for _ in header_cells) + ' |'
        md_lines.append(separator)
        for row in table.rows[1:]:
            row_cells = row.cells
            row_texts = [get_full_text(cell.paragraphs[0]) for cell in row_cells]
            row_text = '| ' + ' | '.join(row_texts) + ' |'
            md_lines.append(row_text)
        md_lines.append('')
    
    md_content = '\n'.join(md_lines)
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    print(f'转换完成: {md_path}')

if __name__ == '__main__':
    files_to_convert = [
        ('/Users/dada/Documents/GitHub/Project-Ark/Docs/GDD/示巴星 (鸣钟者)/示巴星 V1.docx',
         '/Users/dada/Documents/GitHub/Project-Ark/Docs/GDD/示巴星 (鸣钟者)/示巴星 V1.md'),
        ('/Users/dada/Documents/GitHub/Project-Ark/Docs/GDD/示巴星 (鸣钟者)/示巴星 V2.docx',
         '/Users/dada/Documents/GitHub/Project-Ark/Docs/GDD/示巴星 (鸣钟者)/示巴星 V2.md'),
    ]
    
    for docx_path, md_path in files_to_convert:
        if os.path.exists(docx_path):
            docx_to_markdown(docx_path, md_path)
        else:
            print(f'文件不存在: {docx_path}')
